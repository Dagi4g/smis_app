# Copyright 2025 Dagim Genene
# 
# Licensed under the Apache License, Version 2.0 (the "License");
# you may not use this file except in compliance with the License.
# You may obtain a copy of the License at
# 
#     https://www.apache.org/licenses/LICENSE-2.0
# 
# Unless required by applicable law or agreed to in writing, software
# distributed under the License is distributed on an "AS IS" BASIS,
# WITHOUT WARRANTIES OR CONDITIONS OF ANY KIND, either express or implied.
# See the License for the specific language governing permissions and
# limitations under the License.
import re
from docx import Document
from pprint import pprint

class MarkListProcessor:
    def __init__(self, file_path):
        self.doc = Document(file_path)
        self.grade_section_list = []
        self.aranged_dictionary = {}

    def fix_and_extract(self, text):
        """Fix common spelling mistakes and extract grade & section."""
        corrections = {
            r'\bYEER\b': 'YEAR',
            r'\bSCETION\b': 'SECTION',
            r'CHENCHA SECONDARY SCHOOL STUDENTS': 'CHENCHA SECONDARY & PERPARATORY SCHOOL STUDENTS'
        }
        for wrong, correct in corrections.items():
            text = re.sub(wrong, correct, text)

        pattern = r'GRADE & SECTION (\d+)\s*(\w+)'
        matches = re.findall(pattern, text)
        return text, matches

    def extract_grade_sections(self):
        """Extract all grades and sections from the document paragraphs."""
        for paragraph in self.doc.paragraphs:
            _, matches = self.fix_and_extract(paragraph.text)
            for grade, section in matches:
                self.grade_section_list.append(f'{grade} {section}')

    def extract_tables(self):
        """Extract all tables into a dictionary with index as keys."""
        table_dict = {}
        for i, table in enumerate(self.doc.tables):
            table_list = []
            for row in table.rows:
                cells = [cell.text.strip() for cell in row.cells[:4] if cell.text.strip()]
                if cells:  # skip empty rows
                    table_list.append(cells)
            table_dict[str(i)] = table_list
        return table_dict

    def arrange_by_grade_section(self):
        """Map each extracted grade & section to its corresponding table."""
        self.extract_grade_sections()
        table_dict = self.extract_tables()
        arranged_dict = {}

        for i, grade_section in enumerate(self.grade_section_list):
            key = str(i)
            if key in table_dict:
                # Remove header row and incomplete rows
                filtered_rows = [
                    row for row in table_dict[key]
                    if row[0] != 'No' and len(row) >= 2 and row[1].strip()
                ]
                arranged_dict[grade_section] = filtered_rows

        self.aranged_dictionary = arranged_dict
        return self.aranged_dictionary

    def pretty_print(self):
        pprint(self.aranged_dictionary)


# =========================
# Example usage
# =========================
if __name__ == "__main__":
    processor = MarkListProcessor("9th mark.docx")
    processor.arrange_by_grade_section()
    processor.pretty_print()
